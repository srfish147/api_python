from docx import Document
from docx.shared import Inches
from io import BytesIO
import requests

def adicionar_cabecalho_com_logo(input_path, output_path, logo_url):
    doc = Document(input_path)

    for section in doc.sections:
        header = section.header

        # Remove cabeçalhos existentes para evitar duplicação
        for paragraph in header.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)

        # Adiciona novo parágrafo e imagem
        paragraph = header.add_paragraph()
        run = paragraph.add_run()

        response = requests.get(logo_url)
        image_stream = BytesIO(response.content)

        run.add_picture(image_stream, width=Inches(1.2))  # ajuste conforme sua logo

    doc.save(output_path)
