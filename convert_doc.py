import os
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert as docx_to_pdf
from pdf2docx import Converter

def adicionar_cabecalho_com_logo_e_numero_pagina(input_path, output_path, logo_url):
    doc = Document(input_path)

    for section in doc.sections:
        section.header_distance = Inches(0.8)
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Inches(6.27))
        table.autofit = False

        # Remove bordas da tabela
        for cell in table._cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                edge = OxmlElement(f'w:{border}')
                edge.set(qn('w:val'), 'nil')
                tc_borders.append(edge)
            tc_pr.append(tc_borders)

        row = table.rows[0]
        cell_logo = row.cells[0]
        cell_page = row.cells[1]

        # Insere logo
        if logo_url:
            try:
                if logo_url.startswith("//"):
                    logo_url = "https:" + logo_url
                response = requests.get(logo_url)
                image_stream = BytesIO(response.content)
                paragraph_logo = cell_logo.paragraphs[0]
                run_logo = paragraph_logo.add_run()
                run_logo.add_picture(image_stream, height=Inches(0.61))
            except Exception as e:
                print(f"Erro ao carregar logo: {e}")

        # Numeração de página
        paragraph_page = cell_page.paragraphs[0]
        paragraph_page.alignment = 2  # Direita
        run_page = paragraph_page.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)

        # Remove rodapé
        for para in section.footer.paragraphs:
            p = para._element
            p.getparent().remove(p)

    # Insere espaçamento após o cabeçalho
    if doc.paragraphs:
        spacer = doc.paragraphs[0].insert_paragraph_before()
        spacer.paragraph_format.space_before = Pt(18)
        spacer.paragraph_format.space_after = Pt(18)
        spacer.add_run("")

    # Salva DOCX temporário com cabeçalho
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    # Caminhos auxiliares
    pdf_path = output_path.replace(".docx", ".pdf")
    final_docx_path = output_path.replace(".docx", "_final.docx")

    # Converte para PDF
    docx_to_pdf(output_path, pdf_path)

    # Converte PDF para novo DOCX compatível com Google Docs
    converter = Converter(pdf_path)
    converter.convert(final_docx_path, start=0, end=None)
    converter.close()

    return final_docx_path